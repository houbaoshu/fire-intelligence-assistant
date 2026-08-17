"""一次性生成 backend/data/templates/ 下的 docxtpl 模板（.docx）。

模板为真实 Word 文件，含 docxtpl/Jinja2 占位（{{ }} 与 {%tr/%}p 循环）。
模板结构调整后重新运行本脚本即可再生：

    .venv/bin/python scripts/generate_templates.py
"""

from pathlib import Path

from docx import Document

TEMPLATE_DIR = Path(__file__).resolve().parents[1] / "data" / "templates"


def _kv(doc: Document, label: str, var: str) -> None:
    doc.add_paragraph(f"{label}：{{{{ {var} }}}}")


def build_inspection_record() -> Document:
    doc = Document()
    doc.add_heading("{{ title }}", level=1)
    _kv(doc, "记录编号", "record_number")
    _kv(doc, "被检查单位", "inspection_unit")
    _kv(doc, "检查地址", "inspection_address")
    _kv(doc, "检查日期", "inspection_date")
    _kv(doc, "检查人员", "inspector_names")
    _kv(doc, "联系人", "contact_person")
    _kv(doc, "联系电话", "contact_phone")

    doc.add_heading("检查发现问题", level=2)
    # docxtpl 0.20 的 {%tr %} 标签在同一行内会被贪婪匹配吞掉，
    # 因此 for / endfor 各占一个独占行（渲染时整行被替换消失），
    # 中间行才是实际的数据行模板。
    table = doc.add_table(rows=4, cols=6)
    table.style = "Table Grid"
    for cell, header in zip(
        table.rows[0].cells, ["序号", "类型", "位置", "问题描述", "法律依据", "整改要求"]
    ):
        cell.text = header
    loop_start = table.rows[1]
    loop_start.cells[0].merge(loop_start.cells[5]).text = "{%tr for item in items %}"
    row = table.rows[2].cells
    row[0].text = "{{ item.sort_order }}"
    row[1].text = "{{ item.item_type }}"
    row[2].text = "{{ item.location }}"
    row[3].text = "{{ item.description }}"
    row[4].text = "{{ item.legal_basis }}"
    row[5].text = "{{ item.correction_requirement }}"
    loop_end = table.rows[3]
    loop_end.cells[0].merge(loop_end.cells[5]).text = "{%tr endfor %}"

    doc.add_heading("检查情况概述", level=2)
    doc.add_paragraph("{{ summary }}")
    doc.add_heading("检查结论", level=2)
    doc.add_paragraph("{{ conclusion }}")
    return doc


def build_photo_report() -> Document:
    doc = Document()
    doc.add_heading("{{ title }}", level=1)
    _kv(doc, "被检查单位", "inspection_unit")
    _kv(doc, "检查地址", "inspection_address")
    doc.add_heading("违规情况摘要", level=2)
    doc.add_paragraph("{{ violation_summary }}")
    doc.add_heading("现场照片", level=2)
    doc.add_paragraph("{%p for image in images %}")
    doc.add_paragraph("{{ image.image }}")
    doc.add_paragraph("照片说明：{{ image.caption }}")
    doc.add_paragraph("{%p endfor %}")
    return doc


def build_interview_record() -> Document:
    doc = Document()
    doc.add_heading("{{ title }}", level=1)
    _kv(doc, "被询问人", "interviewee_name")
    _kv(doc, "询问人", "interviewer_names")
    _kv(doc, "地点", "location")
    doc.add_paragraph("时间：{{ started_at }} 至 {{ ended_at }}")

    doc.add_heading("询问内容", level=2)
    doc.add_paragraph("{%p for qa in questions_and_answers %}")
    doc.add_paragraph("问：{{ qa.question }}")
    doc.add_paragraph("答：{{ qa.answer }}")
    doc.add_paragraph("{%p endfor %}")

    doc.add_heading("转写原文", level=2)
    doc.add_paragraph("{{ transcript }}")
    return doc


def main() -> None:
    TEMPLATE_DIR.mkdir(parents=True, exist_ok=True)
    builders = {
        "inspection_record.docx": build_inspection_record,
        "photo_report.docx": build_photo_report,
        "interview_record.docx": build_interview_record,
    }
    for name, builder in builders.items():
        path = TEMPLATE_DIR / name
        builder().save(str(path))
        print(f"written: {path}")


if __name__ == "__main__":
    main()
