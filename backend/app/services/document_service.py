"""Document generation & download service.

Documents are ALWAYS rendered from saved structured business data via Word
templates (python-docx). Frontend never generates Word documents.

Versioning (DATABASE.md generated_documents): finalized documents are never
overwritten; regenerating increments the version.
"""
from __future__ import annotations

import io
import re
import uuid
from datetime import datetime, timezone

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from sqlalchemy import select
from sqlalchemy.orm import Session

from app.core.config import get_settings
from app.core.exceptions import ConflictError, DocumentGenerationError, NotFoundError
from app.core.logging import get_logger
from app.models.document import GeneratedDocument
from app.services.audit_service import AuditService
from app.services.file_service import FileService

logger = get_logger("documents")

DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

_PLACEHOLDER_RE = re.compile(r"\{\{\s*([a-zA-Z0-9_.]+)\s*\}\}")


def _utcnow() -> datetime:
    return datetime.now(timezone.utc)


class DocumentService:
    def __init__(self, db: Session):
        self.db = db
        self.files = FileService(db)
        self.audit = AuditService(db)
        self.settings = get_settings()

    # ---- template helpers ---------------------------------------------------

    def _template_path(self, doc_type: str):
        name = {
            "inspection_record_docx": "inspection_record.docx",
            "photo_report_docx": "photo_report.docx",
            "interview_record_docx": "interview_record.docx",
        }.get(doc_type)
        if name is None:
            raise DocumentGenerationError(f"不支持的文档类型:{doc_type}")
        return self.settings.templates_dir / name

    def _ensure_template(self, doc_type: str):
        path = self._template_path(doc_type)
        if not path.exists():
            _create_default_template(path, doc_type)
        return path

    # ---- generation ---------------------------------------------------------

    def generate_document(self, user, entity_type: str, entity_id: uuid.UUID | str) -> GeneratedDocument:
        """Render a Word document from the saved structured record.

        Returns the new GeneratedDocument record (versioned). Raises
        DocumentGenerationError / NotFoundError on failure.
        """
        entity_id = uuid.UUID(str(entity_id))
        doc_type = {
            "inspection_record": "inspection_record_docx",
            "photo_report": "photo_report_docx",
            "interview_record": "interview_record_docx",
        }.get(entity_type)
        if doc_type is None:
            raise DocumentGenerationError(f"未知业务实体:{entity_type}")

        data = self._collect_entity_data(entity_type, entity_id)

        def _image_resolver(uploaded_file_id: str):
            try:
                return self.files.read(uuid.UUID(uploaded_file_id))
            except Exception:  # noqa: BLE001
                logger.warning("image read failed: %s", uploaded_file_id)
                return None

        template_path = self._ensure_template(doc_type)
        try:
            docx_bytes = render_docx(template_path, doc_type, data, image_resolver=_image_resolver)
        except DocumentGenerationError:
            raise
        except Exception as exc:  # noqa: BLE001
            logger.exception("docx render failed for %s %s", entity_type, entity_id)
            raise DocumentGenerationError(f"文书渲染失败:{exc}") from exc

        uploaded = self.files.store_bytes(
            docx_bytes,
            "generated_document",
            f"{doc_type}.docx",
            user.id,
            mime=DOCX_MIME,
        )

        # versioning: next version number
        prev = self.db.scalars(
            select(GeneratedDocument)
            .where(
                GeneratedDocument.source_entity_type == entity_type,
                GeneratedDocument.source_entity_id == entity_id,
            )
            .order_by(GeneratedDocument.version.desc())
        ).first()
        version = (prev.version + 1) if prev else 1

        generated = GeneratedDocument(
            document_type=doc_type,
            source_entity_type=entity_type,
            source_entity_id=entity_id,
            uploaded_file_id=uploaded.id,
            version=version,
            created_by=user.id,
        )
        self.db.add(generated)
        self.audit.log(
            "document.generate", user_id=user.id,
            entity_type=entity_type, entity_id=entity_id,
        )
        self.db.commit()
        return generated

    def _collect_entity_data(self, entity_type: str, entity_id: uuid.UUID) -> dict:
        if entity_type == "inspection_record":
            from app.models.inspection import InspectionRecord

            rec = self.db.get(InspectionRecord, entity_id)
            if rec is None or rec.deleted_at is not None:
                raise NotFoundError("检查记录不存在")
            return {
                "record_number": rec.record_number or "",
                "title": rec.title or "",
                "inspection_unit": rec.inspection_unit or "",
                "inspection_address": rec.inspection_address or "",
                "inspection_date": (rec.inspection_date.isoformat() if rec.inspection_date else ""),
                "inspector_names": "、".join(rec.inspector_names or []),
                "contact_person": rec.contact_person or "",
                "contact_phone": rec.contact_phone or "",
                "summary": rec.summary or "",
                "conclusion": rec.conclusion or "",
                "items": [
                    {
                        "location": it.location or "",
                        "description": it.description,
                        "legal_basis": it.legal_basis or "",
                        "correction_requirement": it.correction_requirement or "",
                        "severity": it.severity or "",
                        "item_type": it.item_type,
                    }
                    for it in rec.items
                ],
            }
        if entity_type == "photo_report":
            from app.models.photo_report import PhotoReport

            rec = self.db.get(PhotoReport, entity_id)
            if rec is None or rec.deleted_at is not None:
                raise NotFoundError("拍照报告不存在")
            return {
                "title": rec.title or "",
                "inspection_unit": rec.inspection_unit or "",
                "inspection_address": rec.inspection_address or "",
                "violation_summary": rec.violation_summary or "",
                "images": [
                    {
                        "caption": img.caption or "",
                        "detected_address": img.detected_address or "",
                        "detected_violation": img.detected_violation or "",
                        "frame_timestamp": img.frame_timestamp,
                        "uploaded_file_id": str(img.uploaded_file_id),
                    }
                    for img in rec.images
                    if img.is_selected
                ],
            }
        if entity_type == "interview_record":
            from app.models.interview import InterviewRecord

            rec = self.db.get(InterviewRecord, entity_id)
            if rec is None or rec.deleted_at is not None:
                raise NotFoundError("询问记录不存在")
            qa = (rec.structured_content or {}).get("questions_and_answers", [])
            return {
                "title": rec.title or "",
                "interviewee_name": rec.interviewee_name or "",
                "interviewer_names": "、".join(rec.interviewer_names or []),
                "location": rec.location or "",
                "started_at": (rec.started_at.isoformat() if rec.started_at else ""),
                "ended_at": (rec.ended_at.isoformat() if rec.ended_at else ""),
                "transcript": rec.transcript or "",
                "questions_and_answers": qa,
            }
        raise DocumentGenerationError(f"未知业务实体:{entity_type}")

    # ---- download -----------------------------------------------------------

    def download_latest(
        self,
        user,
        *,
        entity_type: str,
        entity_id: uuid.UUID,
        fallback_filename: str,
    ) -> tuple[bytes, str]:
        doc = self.db.scalars(
            select(GeneratedDocument)
            .where(
                GeneratedDocument.source_entity_type == entity_type,
                GeneratedDocument.source_entity_id == entity_id,
            )
            .order_by(GeneratedDocument.version.desc())
        ).first()
        if doc is None:
            raise ConflictError("文书尚未生成,请先定稿并等待文书生成完成")
        data = self.files.read(doc.uploaded_file_id)
        self.audit.log(
            "document.download", user_id=user.id,
            entity_type=entity_type, entity_id=entity_id,
            details={"version": doc.version},
        )
        self.db.commit()
        return data, fallback_filename


# ---------------------------------------------------------------------------
# Rendering
# ---------------------------------------------------------------------------

def _set_cell_text(cell, text: str) -> None:
    cell.text = text or ""
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(10.5)


def _ensure_table(doc: Document) -> None:
    if doc.tables:
        return
    doc.add_table(rows=1, cols=6)
    doc.tables[0].style = "Table Grid"


def render_docx(
    template_path,
    doc_type: str,
    data: dict,
    image_resolver=None,
) -> bytes:
    """Render a template file with structured data into a docx byte stream.

    image_resolver(uploaded_file_id) -> bytes | None embeds photos for
    photo reports; it is provided by the caller (DocumentService).
    """
    doc = Document(str(template_path))

    # 1) replace {{placeholder}} in paragraphs
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if "{{" not in run.text:
                continue
            run.text = _replace_placeholders(run.text, data)
        full = paragraph.text
        if "{{" in full:
            # multi-run placeholders: rebuild the paragraph text
            new_text = _replace_placeholders(full, data)
            for run in paragraph.runs:
                run.text = ""
            if paragraph.runs:
                paragraph.runs[0].text = new_text

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if "{{" in paragraph.text:
                        paragraph.text = _replace_placeholders(paragraph.text, data)

    # 2) inject dynamic blocks
    if doc_type == "inspection_record_docx" and data.get("items"):
        _append_inspection_items(doc, data["items"])
    elif doc_type == "photo_report_docx" and data.get("images"):
        _append_photo_images(doc, data["images"], image_resolver)
    elif doc_type == "interview_record_docx" and data.get("questions_and_answers"):
        _append_qa(doc, data["questions_and_answers"])

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _replace_placeholders(text: str, data: dict) -> str:
    def repl(match):
        key = match.group(1)
        return str(data.get(key, "") or "")

    return _PLACEHOLDER_RE.sub(repl, text)


def _append_inspection_items(doc: Document, items: list[dict]) -> None:
    doc.add_heading("检查发现", level=2)
    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(["类型", "位置", "问题描述", "法律依据", "整改要求", "严重程度"]):
        _set_cell_text(hdr[i], h)
    for item in items:
        row = table.add_row().cells
        _set_cell_text(row[0], item.get("item_type", ""))
        _set_cell_text(row[1], item.get("location", ""))
        _set_cell_text(row[2], item.get("description", ""))
        _set_cell_text(row[3], item.get("legal_basis", ""))
        _set_cell_text(row[4], item.get("correction_requirement", ""))
        _set_cell_text(row[5], item.get("severity", ""))


def _append_photo_images(doc: Document, images: list[dict], image_resolver=None) -> None:
    doc.add_heading("现场照片", level=2)
    for idx, img in enumerate(images, start=1):
        p = doc.add_paragraph()
        p.add_run(f"照片 {idx}").bold = True
        ts = img.get("frame_timestamp")
        if ts is not None:
            p.add_run(f"(视频时间点 {ts:.1f}s)")
        caption = img.get("caption") or ""
        doc.add_paragraph(caption)
        address = img.get("detected_address") or ""
        if address:
            doc.add_paragraph(f"地址:{address}")
        if image_resolver:
            try:
                file_bytes = image_resolver(img.get("uploaded_file_id", ""))
                if file_bytes:
                    import tempfile
                    from pathlib import Path

                    with tempfile.NamedTemporaryFile(suffix=".jpg", delete=False) as tf:
                        tf.write(file_bytes)
                        tmp_path = Path(tf.name)
                    try:
                        doc.add_picture(str(tmp_path), width=Inches(5.0))
                    finally:
                        tmp_path.unlink(missing_ok=True)
            except Exception as exc:  # noqa: BLE001
                logger.warning("photo embed failed: %s", exc)
                doc.add_paragraph("[图片无法嵌入]")


def _append_qa(doc: Document, qa: list[dict]) -> None:
    doc.add_heading("询问内容", level=2)
    for idx, item in enumerate(qa, start=1):
        doc.add_paragraph(f"{idx}. 问:{item.get('question', '')}")
        doc.add_paragraph(f"   答:{item.get('answer', '')}")


def _create_default_template(path, doc_type: str) -> None:
    """Bootstrap a minimal real Word template when missing."""
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    title_map = {
        "inspection_record_docx": "消防监督检查记录",
        "photo_report_docx": "消防检查拍照报告",
        "interview_record_docx": "询问笔录",
    }
    title = title_map.get(doc_type, "文书")
    doc.add_heading(title, level=0).alignment = WD_ALIGN_PARAGRAPH.CENTER

    if doc_type == "inspection_record_docx":
        fields = [
            "record_number", "title", "inspection_unit", "inspection_address",
            "inspection_date", "inspector_names", "contact_person", "contact_phone",
            "summary", "conclusion",
        ]
        labels = {
            "record_number": "记录编号", "title": "标题", "inspection_unit": "被检查单位",
            "inspection_address": "检查地址", "inspection_date": "检查日期",
            "inspector_names": "检查人员", "contact_person": "联系人",
            "contact_phone": "联系电话", "summary": "检查情况概述", "conclusion": "检查结论",
        }
        for f in fields:
            doc.add_paragraph(f"{labels[f]}:{{{{ {f} }}}}")
    elif doc_type == "photo_report_docx":
        for f, label in [
            ("title", "标题"), ("inspection_unit", "被检查单位"),
            ("inspection_address", "检查地址"), ("violation_summary", "违规情况摘要"),
        ]:
            doc.add_paragraph(f"{label}:{{{{ {f} }}}}")
    elif doc_type == "interview_record_docx":
        for f, label in [
            ("title", "标题"), ("interviewee_name", "被询问人"),
            ("interviewer_names", "询问人"), ("location", "地点"),
            ("started_at", "开始时间"), ("ended_at", "结束时间"),
        ]:
            doc.add_paragraph(f"{label}:{{{{ {f} }}}}")

    doc.save(str(path))
