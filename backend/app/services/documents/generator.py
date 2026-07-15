"""Word document generation service."""

from __future__ import annotations

import tempfile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.core.logging import get_logger
from app.models.inspection import InspectionRecord
from app.models.interview import InterviewRecord
from app.models.photo_report import PhotoReport

logger = get_logger(__name__)


class DocumentGeneratorService:
    """Service for generating Word documents from structured data.

    Uses python-docx to create professional documents from
    inspection records, photo reports, and interview records.
    """

    async def generate_inspection_record_docx(self, record: InspectionRecord) -> str:
        """Generate a Word document from an inspection record.

        Args:
            record: InspectionRecord model instance with items loaded.

        Returns:
            Path to the generated temporary file.
        """
        logger.info("Generating inspection record document", extra={"record_id": record.id})

        doc = Document()

        # Title
        title = doc.add_heading(record.title or "消防检查记录", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Metadata section
        doc.add_paragraph()
        if record.record_number:
            doc.add_paragraph(f"记录编号: {record.record_number}")
        if record.inspection_unit:
            doc.add_paragraph(f"检查单位: {record.inspection_unit}")
        if record.inspection_address:
            doc.add_paragraph(f"检查地址: {record.inspection_address}")
        if record.inspection_date:
            doc.add_paragraph(f"检查日期: {record.inspection_date.strftime('%Y年%m月%d日')}")
        if record.inspector_names:
            inspectors = (
                ", ".join(record.inspector_names)
                if isinstance(record.inspector_names, list)
                else str(record.inspector_names)
            )
            doc.add_paragraph(f"检查人员: {inspectors}")
        if record.contact_person:
            doc.add_paragraph(f"联系人: {record.contact_person}")
        if record.contact_phone:
            doc.add_paragraph(f"联系电话: {record.contact_phone}")

        doc.add_paragraph()

        # Summary
        if record.summary:
            doc.add_heading("检查概况", level=1)
            doc.add_paragraph(record.summary)

        # Items
        if record.items:
            doc.add_heading("检查项目", level=1)
            for idx, item in enumerate(record.items, 1):
                doc.add_heading(f"{idx}. {item.item_type}", level=2)
                if item.location:
                    doc.add_paragraph(f"位置: {item.location}")
                doc.add_paragraph(f"描述: {item.description}")
                if item.legal_basis:
                    doc.add_paragraph(f"法律依据: {item.legal_basis}")
                if item.correction_requirement:
                    doc.add_paragraph(f"整改要求: {item.correction_requirement}")
                if item.severity:
                    doc.add_paragraph(f"严重程度: {item.severity}")
                doc.add_paragraph()

        # Conclusion
        if record.conclusion:
            doc.add_heading("检查结论", level=1)
            doc.add_paragraph(record.conclusion)

        # Save to temporary file
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            doc.save(tmp.name)
            return tmp.name

    async def generate_photo_report_docx(self, report: PhotoReport) -> str:
        """Generate a Word document from a photo report.

        Args:
            report: PhotoReport model instance with images loaded.

        Returns:
            Path to the generated temporary file.
        """
        logger.info("Generating photo report document", extra={"report_id": report.id})

        doc = Document()

        # Title
        title = doc.add_heading(report.title or "消防照片报告", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Metadata
        doc.add_paragraph()
        if report.inspection_unit:
            doc.add_paragraph(f"检查单位: {report.inspection_unit}")
        if report.inspection_address:
            doc.add_paragraph(f"检查地址: {report.inspection_address}")

        doc.add_paragraph()

        # Violation summary
        if report.violation_summary:
            doc.add_heading("违规情况概述", level=1)
            doc.add_paragraph(report.violation_summary)

        # Images
        if report.images:
            doc.add_heading("现场照片", level=1)
            selected_images = [img for img in report.images if img.is_selected]
            for idx, img in enumerate(selected_images, 1):
                doc.add_heading(f"照片 {idx}", level=2)
                if img.caption:
                    doc.add_paragraph(f"说明: {img.caption}")
                if img.detected_address:
                    doc.add_paragraph(f"地址: {img.detected_address}")
                if img.detected_violation:
                    doc.add_paragraph(f"违规情况: {img.detected_violation}")
                doc.add_paragraph()

        # Save to temporary file
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            doc.save(tmp.name)
            return tmp.name

    async def generate_interview_record_docx(self, record: InterviewRecord) -> str:
        """Generate a Word document from an interview record.

        Args:
            record: InterviewRecord model instance.

        Returns:
            Path to the generated temporary file.
        """
        logger.info("Generating interview record document", extra={"record_id": record.id})

        doc = Document()

        # Title
        title = doc.add_heading(record.title or "询问笔录", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Metadata
        doc.add_paragraph()
        if record.interviewee_name:
            doc.add_paragraph(f"被询问人: {record.interviewee_name}")
        if record.interviewer_names:
            interviewers = (
                ", ".join(record.interviewer_names)
                if isinstance(record.interviewer_names, list)
                else str(record.interviewer_names)
            )
            doc.add_paragraph(f"询问人: {interviewers}")
        if record.location:
            doc.add_paragraph(f"地点: {record.location}")
        if record.started_at:
            doc.add_paragraph(f"开始时间: {record.started_at.strftime('%Y年%m月%d日 %H:%M')}")
        if record.ended_at:
            doc.add_paragraph(f"结束时间: {record.ended_at.strftime('%Y年%m月%d日 %H:%M')}")

        doc.add_paragraph()

        # Transcript
        if record.transcript:
            doc.add_heading("询问内容", level=1)
            doc.add_paragraph(record.transcript)

        # Structured content
        if record.structured_content:
            doc.add_heading("结构化内容", level=1)
            # Convert structured content to readable format
            if isinstance(record.structured_content, dict):
                for key, value in record.structured_content.items():
                    doc.add_paragraph(f"{key}: {value}")

        # Save to temporary file
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            doc.save(tmp.name)
            return tmp.name
