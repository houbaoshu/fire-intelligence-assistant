from __future__ import annotations

import uuid
from io import BytesIO
from typing import Any

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from app.db.models import (
    InspectionRecord,
    InspectionRecordItem,
    InterviewRecord,
    PhotoReport,
    PhotoReportImage,
    UploadedFile,
)
from app.services.storage import StorageProvider


def _base_document(title: str) -> DocumentObject:
    document = Document()
    style = document.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(11)
    heading = document.add_heading(title, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return document


def _field(document: DocumentObject, label: str, value: Any) -> None:
    paragraph = document.add_paragraph()
    paragraph.add_run(f"{label}：").bold = True
    if isinstance(value, list):
        paragraph.add_run("、".join(str(item) for item in value))
    else:
        paragraph.add_run("" if value is None else str(value))


def inspection_docx(record: InspectionRecord, items: list[InspectionRecordItem]) -> bytes:
    document = _base_document(record.title or "消防监督检查记录")
    for label, value in (
        ("记录编号", record.record_number),
        ("被检查单位", record.inspection_unit),
        ("检查地址", record.inspection_address),
        ("检查时间", record.inspection_date),
        ("检查人员", record.inspector_names or []),
        ("联系人", record.contact_person),
        ("联系电话", record.contact_phone),
    ):
        _field(document, label, value)
    document.add_heading("检查发现", level=1)
    if items:
        table = document.add_table(rows=1, cols=6)
        headers = ("序号", "类型", "位置", "事实描述", "法律依据", "整改要求")
        for cell, header in zip(table.rows[0].cells, headers, strict=True):
            cell.text = header
        for index, item in enumerate(items, 1):
            cells = table.add_row().cells
            row_values = (
                index,
                item.item_type,
                item.location,
                item.description,
                item.legal_basis,
                item.correction_requirement,
            )
            for cell, row_value in zip(cells, row_values, strict=True):
                cell.text = "" if row_value is None else str(row_value)
    else:
        document.add_paragraph("未记录检查发现。")
    document.add_heading("检查摘要", level=1)
    document.add_paragraph(record.summary or "")
    document.add_heading("检查结论", level=1)
    document.add_paragraph(record.conclusion or "")
    return _save(document)


def photo_report_docx(
    report: PhotoReport,
    images: list[PhotoReportImage],
    files: dict[uuid.UUID, UploadedFile],
    storage: StorageProvider,
) -> bytes:
    document = _base_document(report.title or "消防检查图像报告")
    _field(document, "被检查单位", report.inspection_unit)
    _field(document, "检查地址", report.inspection_address)
    _field(document, "问题摘要", report.violation_summary)
    selected = sorted(
        (image for image in images if image.is_selected), key=lambda item: item.sort_order
    )
    for index, image in enumerate(selected, 1):
        document.add_heading(f"照片 {index}", level=1)
        uploaded = files[image.uploaded_file_id]
        with storage.open(uploaded.storage_path) as source:
            document.add_picture(source, width=Cm(15.5))
        _field(document, "视频时间", image.frame_timestamp)
        _field(document, "图片说明", image.caption)
        _field(document, "识别地址", image.detected_address)
        _field(document, "问题描述", image.detected_violation)
    return _save(document)


def interview_docx(record: InterviewRecord) -> bytes:
    document = _base_document(record.title or "消防安全询问笔录")
    for label, value in (
        ("被询问人", record.interviewee_name),
        ("询问人员", record.interviewer_names or []),
        ("询问地点", record.location),
        ("开始时间", record.started_at),
        ("结束时间", record.ended_at),
    ):
        _field(document, label, value)
    document.add_heading("结构化笔录", level=1)
    content = record.structured_content or {}
    sections = content.get("sections") if isinstance(content, dict) else None
    if isinstance(sections, list):
        for section in sections:
            if isinstance(section, dict):
                _field(document, str(section.get("question") or "问题"), section.get("answer"))
    else:
        document.add_paragraph(str(content))
    return _save(document)


def _save(document: DocumentObject) -> bytes:
    output = BytesIO()
    document.save(output)
    return output.getvalue()
