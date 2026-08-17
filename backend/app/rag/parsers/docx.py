"""DOCX 解析（python-docx）：按段落输出文本块，页码不可得（None）。"""

import io

from app.rag.parsers.base import ParsedBlock, ParsedDocument, parse_error


def parse_docx(data: bytes) -> ParsedDocument:
    try:
        import docx
    except ImportError:
        raise parse_error("DOCX 解析组件未安装，请联系管理员")
    try:
        document = docx.Document(io.BytesIO(data))
        blocks = [
            ParsedBlock(text=p.text.strip())
            for p in document.paragraphs
            if p.text.strip()
        ]
        for table in document.tables:
            for row in table.rows:
                line = " ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if line:
                    blocks.append(ParsedBlock(text=line))
    except Exception as exc:
        from app.core.exceptions import AppException

        if isinstance(exc, AppException):
            raise
        raise parse_error("DOCX 文件损坏或格式不受支持，无法解析")
    return ParsedDocument(blocks=blocks)
