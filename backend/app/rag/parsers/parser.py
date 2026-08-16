"""Document parsing: extract normalized text from uploaded source documents."""
from __future__ import annotations

import io
from abc import ABC, abstractmethod
from pathlib import Path

from app.core.exceptions import ValidationError


class ParseResult:
    def __init__(self, text: str, pages: list[str] | None = None):
        self.text = text
        self.pages = pages  # page-wise text (best effort)


class DocumentParser(ABC):
    extensions: tuple[str, ...] = ()

    @abstractmethod
    def parse_bytes(self, data: bytes) -> ParseResult:
        ...


class PdfParser(DocumentParser):
    extensions = (".pdf",)

    def parse_bytes(self, data: bytes) -> ParseResult:
        from pypdf import PdfReader

        try:
            reader = PdfReader(io.BytesIO(data))
        except Exception as e:
            raise ValidationError(f"PDF 解析失败:{e}") from e
        pages = []
        for page in reader.pages:
            try:
                pages.append(page.extract_text() or "")
            except Exception:
                pages.append("")
        return ParseResult(text="\n\n".join(pages), pages=pages)


class DocxParser(DocumentParser):
    extensions = (".docx",)

    def parse_bytes(self, data: bytes) -> ParseResult:
        from docx import Document

        try:
            doc = Document(io.BytesIO(data))
        except Exception as e:
            raise ValidationError(f"DOCX 解析失败:{e}") from e
        paragraphs = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
        # include table text
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                paragraphs.append(" | ".join(cells))
        return ParseResult(text="\n".join(paragraphs))


class PlainTextParser(DocumentParser):
    extensions = (".txt", ".md")

    def parse_bytes(self, data: bytes) -> ParseResult:
        try:
            text = data.decode("utf-8")
        except UnicodeDecodeError:
            try:
                text = data.decode("gbk")
            except UnicodeDecodeError:
                raise ValidationError("文本文件编码不支持,请使用 UTF-8") from None
        return ParseResult(text=text)


# .doc / .ppt / .pptx are accepted for upload but need conversion tooling;
# we report a clear, actionable error rather than silently skipping content.
class UnsupportedForParsingParser(DocumentParser):
    extensions = (".doc", ".ppt", ".pptx")

    def parse_bytes(self, data: bytes) -> ParseResult:
        raise ValidationError(
            "该格式(.doc/.ppt/.pptx)暂不支持直接解析,请转换为 PDF 或 DOCX 后重新上传"
        )


_PARSERS: list[DocumentParser] = [
    PdfParser(),
    DocxParser(),
    PlainTextParser(),
    UnsupportedForParsingParser(),
]
_PARSER_BY_EXT = {ext: p for p in _PARSERS for ext in p.extensions}


def parse_document(data: bytes, extension: str) -> ParseResult:
    ext = extension.lower()
    parser = _PARSER_BY_EXT.get(ext)
    if parser is None:
        raise ValidationError(f"不支持的文档格式:{ext or '未知'}")
    result = parser.parse_bytes(data)
    if not result.text.strip():
        raise ValidationError("文档中没有可解析的文本内容(可能为扫描件,当前版本不支持扫描件 OCR)")
    return result
